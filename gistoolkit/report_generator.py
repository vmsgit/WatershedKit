from pathlib import Path
from datetime import datetime
import math
import subprocess
import shutil

import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# GENERAL HELPERS
# ============================================================

def clean_value(value):
    """Convert NumPy values to normal Python values."""
    if isinstance(value, np.generic):
        return value.item()

    if isinstance(value, float) and math.isnan(value):
        return None

    return value


def format_number(value, decimals=2):
    value = clean_value(value)

    if value is None:
        return "—"

    if isinstance(value, (int, np.integer)):
        return f"{value:,}"

    if isinstance(value, (float, int)):
        return f"{value:,.{decimals}f}"

    return str(value)


def format_percent(value, decimals=2):
    value = clean_value(value)

    if value is None:
        return "—"

    return f"{value:,.{decimals}f}%"


def format_km_from_m(value, decimals=2):
    value = clean_value(value)
    if value is None:
        return "—"
    return f"{value / 1000:,.{decimals}f} km"


def format_area_km2(value, decimals=2):
    value = clean_value(value)
    if value is None:
        return "—"
    return f"{value / 1_000_000:,.{decimals}f} km²"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="D9EAF7"):
    """Apply background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = tblPr.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7B7B7")


def add_parameter_table(doc, rows):
    """
    rows = [
        ("Parameter Name", "Value", "Unit"),
        ...
    ]
    """
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Parameter", "Value", "Unit"]

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i])

    for name, value, unit in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], value)
        set_cell_text(cells[2], unit)

    add_table_borders(table)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_map(doc, map_path, title, caption=None, width=6.2):
    """
    Insert a PNG map into the report.
    """
    if not map_path:
        return

    path = Path(map_path)

    if not path.exists():
        doc.add_paragraph(f"[Map not found: {path}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    caption_text = caption or title

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(9)


# ============================================================
# DOCUMENT SETUP
# ============================================================

def setup_document(doc):
    section = doc.sections[0]

    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles

    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True

    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True

    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True


# ============================================================
# COVER PAGE
# ============================================================

def add_cover_page(doc, results):
    project_name = results.get("project_name", "Watershed Analysis")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run("\n\n\n")

    title = p.add_run("WATERSHED / CATCHMENT\nGIS ANALYSIS REPORT")
    title.bold = True
    title.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"\nProject: {project_name}")
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(
        f"\nGenerated: {datetime.now().strftime('%d %B %Y')}"
    )
    run.font.size = Pt(10)

    doc.add_page_break()


# ============================================================
# SECTION 1
# ============================================================

def add_project_information(doc, results):
    add_heading(doc, "1. Project and Input Information", 1)

    rows = [
        (
            "Project Name",
            results.get("project_name", "—"),
            ""
        ),
        (
            "Project CRS",
            results.get("dem_crs", "—"),
            ""
        ),
        (
            "Stream Extraction Threshold",
            format_number(results.get("streams_threshold"), 0),
            "cells"
        ),
        (
            "User-Selected Contour Interval",
            format_number(
                results.get("user_selected_contour_interval"), 2
            ),
            "m"
        ),
        (
            "Outlet Latitude (WGS84)",
            format_number(
                results.get("outlet_latitude_wgs84"), 6
            ),
            "°"
        ),
        (
            "Outlet Longitude (WGS84)",
            format_number(
                results.get("outlet_longitude_wgs84"), 6
            ),
            "°"
        ),
    ]

    add_parameter_table(doc, rows)


# ============================================================
# SECTION 2
# ============================================================

def add_catchment_overview(doc, results, maps):
    add_heading(doc, "2. Catchment Characteristics", 1)

    rows = [
        (
            "Catchment Area",
            format_area_km2(results.get("catchment_area_m2")),
            "km²"
        ),
        (
            "Catchment Perimeter",
            format_km_from_m(results.get("catchment_perimeter_m")),
            "km"
        ),
        (
            "Basin Length",
            format_km_from_m(results.get("basin_length")),
            "km"
        ),
        (
            "Catchment Centroid Latitude",
            format_number(
                results.get("catchment_centroid_latitude"), 3
            ),
            "project CRS"
        ),
        (
            "Catchment Centroid Longitude",
            format_number(
                results.get("catchment_centroid_longitude"), 3
            ),
            "project CRS"
        ),
        (
            "Length to Centroid",
            format_km_from_m(results.get("length_to_centroid_m")),
            "km"
        ),
        (
            "Minimum Catchment Elevation",
            format_number(results.get("watershed_min_elev"), 2),
            "m"
        ),
        (
            "Maximum Catchment Elevation",
            format_number(results.get("watershed_max_elev"), 2),
            "m"
        ),
        (
            "Basin Relief",
            format_number(results.get("basin_relief_m"), 2),
            "m"
        ),
        (
            "Relief Ratio",
            format_number(results.get("relief_ratio_m"), 4),
            ""
        ),
        (
            "True Average Basin Slope",
            format_percent(results.get("true_basin_slope_percent")),
            "%"
        ),
    ]

    add_parameter_table(doc, rows)

    add_map(
        doc,
        maps.get("centroid"),
        "Catchment Centroid",
        "Figure 1. Catchment centroid and Length to centroid."
    )


# ============================================================
# SECTION 3
# ============================================================

def add_drainage_network(doc, results, maps):
    add_heading(doc, "3. Drainage Network Analysis", 1)

    rows = [
        (
            "Maximum Stream Order",
            format_number(results.get("outlet_stream_order"), 0),
            ""
        ),
        (
            "Number of Stream Segments",
            format_number(results.get("nos_streamsof_catchment"), 0),
            ""
        ),
        (
            "Total Stream Length",
            format_km_from_m(results.get("all_streams_length")),
            "km"
        ),
        (
            "Drainage Density",
            format_number(results.get("drainage_density_km"), 3),
            "km/km²"
        ),
        (
            "Stream Frequency",
            format_number(results.get("stream_frequency_km2"), 3),
            "streams/km²"
        ),
        (
            "Constant of Channel Maintenance",
            format_number(
                results.get("constant_channel_maintenance_km"), 3
            ),
            "km²/km"
        ),
        (
            "Infiltration Number",
            format_number(
                results.get("infiltration_number_km"), 3
            ),
            ""
        ),
    ]

    add_parameter_table(doc, rows)

    add_heading(doc, "Stream-Order Statistics", 2)

    summary = results.get("streams_summary", {})

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Stream Order",
        "Number of Streams",
        "Total Length (m)",
        "Mean Length (m)",
        "Bifurcation Ratio (Rb)",
        "Stream Length Ratio (Rl)",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i])

    for order, data in sorted(summary.items()):
        cells = table.add_row().cells

        values = [
            format_number(order, 0),
            format_number(data.get("n_streams"), 0),
            format_number(data.get("total_length_m"), 2),
            format_number(data.get("mean_length_m"), 2),
            format_number(data.get("Rb"), 3),
            format_number(data.get("Rl"), 3),
        ]

        for i, value in enumerate(values):
            set_cell_text(cells[i], value)

    add_table_borders(table)
    doc.add_paragraph()

    add_map(
        doc,
        maps.get("stream_network"),
        "Stream Network",
        "Figure 2. Stream network of catchment."
    )


# ============================================================
# SECTION 4
# ============================================================

def add_longest_flow_path(doc, results, maps):
    add_heading(doc, "4. Longest Flow Path Analysis", 1)

    p = doc.add_paragraph()
    p.add_run("Definition: ").bold = True
    p.add_run(
        "The longest flow path represents the complete hydrologic "
        "flow path from the most distant point of the catchment to "
        "the outlet. It includes both overland flow and stream/channel flow."
    )

    rows = [
        (
            "Longest Flow Path Length",
            format_km_from_m(results.get("length_longest_path")),
            "km"
        ),
        (
            "Upstream Elevation",
            format_number(
                results.get("longest_path_up_elevation"), 2
            ),
            "m"
        ),
        (
            "Outlet Elevation",
            format_number(
                results.get("longest_path_down_elevation"), 2
            ),
            "m"
        ),
        (
            "Average Longest Flow Path Slope",
            format_percent(
                results.get("longest_path_avgslope_percent")
            ),
            "%"
        ),
        (
            "Time of Concentration — Kirpich",
            format_number(
                results.get("timeof_concentration_kirpich_min"), 2
            ),
            "minutes"
        ),
    ]

    add_parameter_table(doc, rows)

    add_map(
        doc,
        maps.get("longest_flow_path"),
        "Longest Flow Path",
        "Figure 3. Longest flow path from the most distant catchment point to the outlet."
    )


# ============================================================
# SECTION 5
# ============================================================

def add_main_channel(doc, results, slope_results, maps):
    add_heading(doc, "5. Main Channel Analysis", 1)

    p = doc.add_paragraph()
    p.add_run("Definition: ").bold = True
    p.add_run(
        "The main channel is the longest defined stream path from "
        "the upstream stream head to the outlet. Unlike the longest "
        "flow path, it does not include the overland-flow portion "
        "between the catchment boundary and the stream head."
    )

    rows = [
        (
            "Main Channel Length",
            format_km_from_m(results.get("main_channel_length")),
            "km"
        ),
        (
            "Main Channel Outlet Elevation",
            format_number(
                results.get("main_outlet_elevation"), 2
            ),
            "m"
        ),
        (
            "Main Channel Upstream Elevation",
            format_number(
                results.get("main_upstream_elevation"), 2
            ),
            "m"
        ),
        (
            "Main Channel Slope",
            format_percent(
                results.get("main_channel_slope_percent")
            ),
            "%"
        ),
        (
            "Main Channel Outlet X",
            format_number(
                results.get("main_channel_outlet_x"), 2
            ),
            "Project CRS"
        ),
        (
            "Main Channel Outlet Y",
            format_number(
                results.get("main_channel_outlet_y"), 2
            ),
            "Project CRS"
        ),
        (
            "Main Channel Upstream X",
            format_number(
                results.get("main_channel_upstream_x"), 2
            ),
            "Project CRS"
        ),
        (
            "Main Channel Upstream Y",
            format_number(
                results.get("main_channel_upstream_y"), 2
            ),
            "Project CRS"
        ),
    ]

    add_parameter_table(doc, rows)

    add_heading(doc, "Longest Flow Path vs Main Channel", 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headers = ["Characteristic", "Longest Flow Path", "Main Channel"]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i])

    comparison = [
        (
            "Overland flow",
            "Included",
            "Not included"
        ),
        (
            "Defined stream flow",
            "Included",
            "Included"
        ),
        (
            "Starting point",
            "Most distant catchment point",
            "Upstream stream head"
        ),
        (
            "Ending point",
            "Outlet",
            "Outlet"
        ),
    ]

    for row in comparison:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    add_table_borders(table)
    doc.add_paragraph()

    add_map(
        doc,
        maps.get("main_channel"),
        "Main Channel",
        "Figure 4. Main channel from stream head to outlet."
    )


# ============================================================
# SECTION 6
# ============================================================

def add_channel_slope_analysis(doc, slope_results):
    add_heading(doc, "6. Main Channel Slope Analysis", 1)

    doc.add_paragraph(
        "The main channel slope is evaluated using multiple approaches. "
        "The simple H/L slope provides an end-to-end slope, while the "
        "equivalent-slope approaches account for variations in channel "
        "slope along the channel profile."
    )

    simple = slope_results.get("simple_slope", {})
    equal_elev = slope_results.get("equal_elevation_drop") or {
    "Sc_m_per_km": None,
    "Sc_fraction": None,
    "n_segments": None,
    "converged": False,
}
    equal_len = slope_results.get("equal_length", {})

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = [
        "Method",
        "Slope (m/km)",
        "Slope (%)",
        "Segments",
        "Status",
    ]

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i])

    rows = [
        (
            "Simple End-to-End Slope (H/L)",
            format_number(simple.get("slope_m_per_km"), 2),
            format_percent(
                simple.get("slope_fraction", None) * 100
                if simple.get("slope_fraction") is not None
                else None
            ),
            "—",
            "Calculated",
        ),
        (
            "Equal Elevation-Drop Equivalent Slope",
            format_number(equal_elev.get("Sc_m_per_km"), 2),
            format_percent(
                equal_elev.get("Sc_fraction", None) * 100
                if equal_elev.get("Sc_fraction") is not None
                else None
            ),
            format_number(equal_elev.get("n_segments"), 0),
            "Converged" if equal_elev.get("converged") else "Not converged",
        ),
        (
            "Equal Length Equivalent Slope",
            format_number(equal_len.get("Sc_m_per_km"), 2),
            format_percent(
                equal_len.get("Sc_fraction", None) * 100
                if equal_len.get("Sc_fraction") is not None
                else None
            ),
            format_number(equal_len.get("n_segments"), 0),
            "Converged" if equal_len.get("converged") else "Not converged",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    add_table_borders(table)
    doc.add_paragraph()

    method_used = slope_results.get("method_used")

    p = doc.add_paragraph()
    p.add_run("Method applied by calculation routine: ").bold = True
    p.add_run(str(method_used or "—"))

    # Technical information
    add_heading(doc, "Channel Slope Calculation Information", 2)

    rows = [
        (
            "Catchment Area",
            format_number(
                slope_results.get("catchment_area_km2"), 2
            ),
            "km²"
        ),
        (
            "Channel Length",
            format_number(
                slope_results.get("channel_length_km"), 2
            ),
            "km"
        ),
        (
            "DEM Cell Size",
            format_number(
                slope_results.get("dem_cell_size_m"), 2
            ),
            "m"
        ),
        (
            "Elevation Sample Points",
            format_number(
                slope_results.get("sample_points"), 0
            ),
            ""
        ),
        (
            "Estimated DEM Noise",
            format_number(
                slope_results.get("dem_noise_sigma_m"), 3
            ),
            "m"
        ),
        (
            "Total Relief Along Main Channel",
            format_number(
                slope_results.get("total_relief_m"), 2
            ),
            "m"
        ),
    ]

    add_parameter_table(doc, rows)


# ============================================================
# SECTION 7
# ============================================================

def add_morphometric_analysis(doc, results):
    add_heading(doc, "7. Morphometric Analysis", 1)

    rows = [
        (
            "Form Factor (F)",
            format_number(results.get("form_factor"), 4),
            ""
        ),
        (
            "Elongation Ratio (Re)",
            format_number(results.get("elongation_ratio"), 4),
            ""
        ),
        (
            "Circulatory Ratio (Rc)",
            format_number(results.get("circulatory_ratio"), 4),
            ""
        ),
        (
            "Compactness Coefficient (Cc)",
            format_number(
                results.get("compactness_coefficient"), 4
            ),
            ""
        ),
        (
            "Constant of Channel Maintenance (C)",
            format_number(
                results.get("constant_channel_maintenance_km"),
                4
            ),
            "km²/km"
        ),
        (
            "Infiltration Number (If)",
            format_number(
                results.get("infiltration_number_km"), 4
            ),
            ""
        ),
        (
            "Relief Ratio (Rh)",
            format_number(
                results.get("relief_ratio_m"), 4
            ),
            ""
        ),
        (
            "Ruggedness Number (Rn)",
            format_number(
                results.get("ruggedness_number_km"), 4
            ),
            ""
        ),
        (
            "True Average Basin Slope",
            format_percent(
                results.get("true_basin_slope_percent")
            ),
            "%"
        ),
    ]

    add_parameter_table(doc, rows)


# ============================================================
# SECTION 8
# ============================================================

def add_contour_analysis(doc, results, maps):
    add_heading(doc, "8. Contour Analysis", 1)

    rows = [
        (
            "User-Selected Contour Interval",
            format_number(
                results.get("user_selected_contour_interval"), 2
            ),
            "m"
        )
    ]

    add_parameter_table(doc, rows)

    add_map(
        doc,
        maps.get("contour_path"),
        "Contour Path",
        "Figure 5. Contour Map for the Catchment."
    )


# ============================================================
# MAP: OUTLET
# ============================================================

def add_outlet_map(doc, maps):
    add_heading(doc, "9. Outlet Location", 1)

    add_map(
        doc,
        maps.get("outlet"),
        "Outlet Location",
        "Figure 6. Outlet location and catchment boundary."
    )

# ============================================================
# MAP: Hypsometric Curve
# ============================================================

def add_hypsometric_map(doc, maps):
    add_heading(doc, "10. Hypsometric Curve", 1)

    add_map(
        doc,
        maps.get("hypsometric_curve"),
        "Hypsometric Curve",
        "Figure 7. Hypsometric Curve for the Catchment."
    )

# ============================================================
# WARNINGS
# ============================================================

def add_warnings(doc, slope_results):
    warnings = slope_results.get("warnings", [])

    if not warnings:
        return

    add_heading(doc, "10. Calculation Notes and Warnings", 1)

    p = doc.add_paragraph(
        "The following notes were generated during the process for Main Channel Equivalent/Avg slope calculations:"
    )
    doc.add_paragraph("Check the threshold value in whiteboxtool's jenson_snap_pour_points if pour point snapping fails/incorrect.")

    for warning in warnings:
        doc.add_paragraph(
            str(warning),
            style="List Bullet"
        )


# ============================================================
# FOOTER
# ============================================================

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("Watershed GIS Analysis Report  |  Page ")

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


# ============================================================
# MAIN REPORT GENERATOR
# ============================================================

def generate_report(
    results,
    slope_results,
    maps,
    output_docx
):
    """
    Generate watershed analysis report.

    Parameters
    ----------
    results : dict
        Main GIS results dictionary.

    slope_results : dict
        Dictionary returned by channel slope calculation.

    maps : dict
        Dictionary containing paths to PNG maps.

        Expected keys:
            outlet
            stream_network
            centroid
            longest_flow_path
            main_channel
            contour_path

    output_docx : str or Path
        Output DOCX path.

    output_pdf : str or Path, optional
        Output PDF path.
        PDF conversion requires LibreOffice.
    """

    output_docx = Path(output_docx)

    output_docx.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    doc = Document()

    setup_document(doc)

    # --------------------------------------------------------
    # Build report
    # --------------------------------------------------------

    add_cover_page(doc, results)

    add_project_information(
        doc,
        results
    )

    add_catchment_overview(
        doc,
        results,
        maps
    )

    add_drainage_network(
        doc,
        results,
        maps
    )

    add_longest_flow_path(
        doc,
        results,
        maps
    )

    add_main_channel(
        doc,
        results,
        slope_results,
        maps
    )

    add_channel_slope_analysis(
        doc,
        slope_results
    )

    add_morphometric_analysis(
        doc,
        results
    )

    add_contour_analysis(
        doc,
        results,
        maps
    )

    add_outlet_map(
        doc,
        maps
    )

    add_hypsometric_map(
        doc,                
        maps
    )

    add_warnings(
        doc,
        slope_results
    )

    add_page_numbers(doc)

    doc.save(output_docx)

    return output_docx