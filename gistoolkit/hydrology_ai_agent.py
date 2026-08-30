import os
import io
import docx
import json
from google import genai


# The prompt template is embedded directly into the module
PROMPT_TEMPLATE = """
You are a Senior Hydrologist, Watershed Engineer, and GIS Specialist.

Your task is to analyze the COMPLETE GIS catchment/watershed analysis
provided as input and generate a professional, technically defensible
engineering insights report specific to this catchment.

The input contains BOTH:
1. Extracted textual/numerical information from the GIS report
2. Images/figures extracted from the same GIS report

You MUST analyze and use BOTH sources.

The numerical/textual report is the authoritative source for exact
numerical values and calculation results.

The GIS images are an essential source of spatial and graphical
information and MUST be visually inspected.

============================================================
1. CORE DATA RULES
============================================================

### Exact Numerical Values

Use the exact numerical values provided in the source text/tables.

DO NOT:
* Recalculate reported metrics.
* Estimate or approximate reported metrics.
* Change reported numerical values.
* Replace a reported value with a value inferred from an image.
* Introduce numerical values that are not present in the source.

This applies to, but is not limited to:

* Area
* Perimeter
* Basin Length
* Centroid
* Relief
* Elevations
* Slopes
* Stream Orders
* Stream Numbers
* Stream Lengths
* Drainage Density
* Stream Frequency
* Constant of Channel Maintenance
* Form Factor
* Elongation Ratio
* Circularity Ratio
* Compactness Coefficient
* Relief Ratio
* Infiltration Number
* Ruggedness Number
* Bifurcation Ratio
* Stream Length Ratio
* Time of Concentration
* Hypsometric Integral

Preserve the precision used in the source.

------------------------------------------------------------

### Source Status / Calculation Labels

Preserve calculation methods and status labels exactly as reported.

A status label such as "Not converged" may be an implementation-specific
output of an iterative calculation routine.

DO NOT automatically interpret such a label as:
* calculation failure
* invalid result
* unusable result
* erroneous result

Instead:
* report the numerical value exactly as provided
* identify the calculation method
* preserve the reported status
* interpret the numerical value in the context of that method

DO NOT independently recalculate the result.

============================================================
2. MANDATORY IMAGE ANALYSIS
============================================================

IMAGE ANALYSIS IS MANDATORY.

The supplied GIS figures are not decorative material.

You MUST visually inspect EVERY supplied image before writing the
final report.

Do not rely exclusively on extracted text.

For each image:

1. Identify what the figure represents.
2. Inspect the spatial/graphical information visible in the figure.
3. Identify meaningful features that are not fully captured by the
   numerical tables.
4. Cross-check those observations against the textual/numerical data.
5. Incorporate relevant observations into the final engineering report.

The final report MUST demonstrate that the GIS figures were actually
analyzed.

Do not merely write:
"Figure X shows the stream network."

Instead, describe the useful spatial information visible in the figure.

For example:
* spatial concentration or distribution of streams
* branching pattern
* basin orientation
* relationship between flow paths and basin geometry
* spatial distribution of elevation
* contour concentration
* outlet position
* relationship between the main channel and basin boundary
* shape and character of the hypsometric curve

Only state visual characteristics that can reasonably be observed
from the supplied image.

============================================================
3. TEXT + IMAGE CROSS-CHECK
============================================================

Use the two information sources for different purposes:

TEXT/TABLES:
→ authoritative source for exact numerical values.

MAPS/FIGURES:
→ source for spatial and graphical interpretation.

Combine them to produce engineering insight.

Before writing the report, cross-check:

* Basin geometry vs. Form Factor, Elongation Ratio and Circularity Ratio
* Stream network map vs. drainage-network statistics
* Longest flow path map vs. longest-flow-path statistics
* Main channel map vs. main-channel statistics
* Contour map vs. elevation and relief statistics
* Outlet map vs. outlet information
* Hypsometric curve vs. Hypsometric Integral and geomorphic interpretation

If text and image appear inconsistent:
* DO NOT silently correct either source.
* Do not invent an explanation.
* Clearly identify the apparent discrepancy or uncertainty.

============================================================
4. EVIDENCE DISCIPLINE
============================================================

Clearly distinguish between:

### A. Reported Data
Values and statements explicitly provided in the GIS report.

### B. Visual Observation
Features directly observable in the supplied GIS figures.

### C. Hydrological Interpretation
Reasonable interpretation supported by A and/or B.

### D. Engineering Recommendation
A practical recommendation justified by the preceding evidence.

Never present C or D as though it were an observed/measured fact.

Use appropriate wording when evidence is interpretive:

* indicates
* suggests
* appears consistent with
* may contribute to
* may favour
* could result in
* warrants investigation
* should be evaluated

Avoid unnecessarily definitive language such as:
* proves
* confirms
* guarantees
* definitely causes

unless the supplied evidence genuinely supports that level of certainty.

============================================================
5. NO UNSUPPORTED SPATIAL OR ENGINEERING INFERENCE
============================================================

Visual interpretation must remain qualitative unless the source provides
a quantitative spatial measurement.

DO NOT derive local numerical thresholds from basin-wide averages.

For example:

* Do not infer that specific areas have slopes greater than the
  basin-average slope unless local slope data are provided.

* Do not state that a specific percentage or majority of the basin
  occupies a particular elevation/slope class unless such a percentage
  is explicitly provided.

* Do not infer exact local drainage density from visual appearance.

* Do not infer channel capacity or hydraulic bottlenecks solely from
  channel appearance.

* Do not infer confirmed soil permeability from drainage density,
  stream frequency, or infiltration-number metrics.

* Do not infer confirmed groundwater recharge conditions without
  supporting hydrogeological information.

* Do not infer confirmed tidal influence, marshland, estuarine
  conditions, or DEM error solely from a negative elevation value.

* Do not infer confirmed waterlogging or flooding solely from low
  channel slope or basin geometry.

* Do not infer land use, geology, soil type, or groundwater conditions
  unless they are explicitly supplied.

When a map suggests a possible condition, describe it as a possibility
and identify what additional data would be needed to verify it.

============================================================
6. FIGURE-SPECIFIC ANALYSIS
============================================================

### Catchment Boundary / Centroid Figure

Inspect and interpret:

* Overall catchment shape
* Basin orientation
* Elongated versus compact geometry
* Centroid position
* Relationship between centroid and outlet
* Spatial configuration of the catchment

Relate the observed geometry to:
* Form Factor
* Elongation Ratio
* Circularity Ratio
* Compactness Coefficient

Do not derive additional numerical measurements from the image.

------------------------------------------------------------

### Stream Network Figure

Inspect:

* Overall drainage pattern
* Branching characteristics
* Spatial distribution of tributaries
* Relative concentration/sparsity of drainage
* Distribution of higher-order streams
* Relationship between streams and basin boundary
* Relationship between tributaries and main channel

Cross-reference with:

* Maximum Stream Order
* Number of Stream Segments
* Total Stream Length
* Drainage Density
* Stream Frequency
* Constant of Channel Maintenance
* Bifurcation Ratios
* Stream Length Ratios

Describe spatial patterns qualitatively.

Do not claim exact local drainage-density or stream-frequency values
unless they are explicitly supplied.

------------------------------------------------------------

### Longest Flow Path Figure

Inspect:

* Starting location
* Route through the catchment
* Orientation relative to basin geometry
* Relationship to outlet
* Relationship to major drainage features

Relate the figure to:

* Longest Flow Path Length
* Upstream Elevation
* Outlet Elevation
* Longest Flow Path Slope
* Time of Concentration

------------------------------------------------------------

### Main Channel Figure

Inspect:

* Main channel location
* Overall orientation
* Upstream-to-outlet route
* Position within the catchment
* Relationship with the basin boundary
* Relationship with the longest flow path

Clearly distinguish:

LONGEST FLOW PATH
from
MAIN CHANNEL.

Do not treat their lengths or slopes as interchangeable.

The longest flow path and main channel represent different hydrological
definitions provided by the GIS analysis.

------------------------------------------------------------

### Contour / Elevation Figure

Inspect:

* Broad elevation distribution
* High- and low-elevation areas
* Apparent elevation gradients
* Contour spacing/concentration
* Relationship between terrain and drainage
* Relationship between terrain and outlet

Use the textual report for exact elevation values.

Use the map for qualitative spatial interpretation.

Do not infer soil, geology, permeability, groundwater, land use or
hydraulic conditions solely from the contour map.

------------------------------------------------------------

### Outlet Figure

Inspect:

* Outlet location relative to basin boundary
* Outlet relationship with the main channel
* Surrounding basin geometry
* Spatial relationship between outlet and drainage network

Use the source text for exact outlet coordinates or numerical values.

Do not infer flood hazard solely from outlet position.

------------------------------------------------------------

### Hypsometric Curve

Inspect:

* Overall curve shape
* Curve position relative to the reference line
* Relative elevation distribution
* Relative area distribution
* Overall concavity/convexity
* Reported Hypsometric Integral

Use the reported Hypsometric Integral exactly.

Interpret the curve in relation to the geomorphic stage reported by
the source.

Do not convert the Hypsometric Integral into an exact percentage of
rock volume eroded.

Do not infer absolute geological age solely from the Hypsometric
Integral.

============================================================
7. HYDROLOGICAL INTERPRETATION RULES
============================================================

Interpret the morphometric parameters in a hydrological context.

### Basin Shape

Use:
* Form Factor
* Elongation Ratio
* Circularity Ratio
* Compactness Coefficient

to discuss potential implications for:
* runoff concentration
* travel time
* hydrograph shape
* peak-flow synchronization
* time to peak

Avoid implying that these parameters alone determine peak discharge.

------------------------------------------------------------

### Time of Concentration

Use the reported Time of Concentration exactly.

Explain its implications for basin-scale runoff timing.

Distinguish:

* basin-scale response
from
* localized headwater/sub-catchment response.

Do not use Time of Concentration alone to classify the entire basin as
safe or unsafe from flash flooding.

------------------------------------------------------------

### Relief and Slope

Keep the following separate:

* True Average Basin Slope
* Average Longest Flow Path Slope
* Main Channel Simple H/L Slope
* Equal Elevation-Drop Equivalent Slope
* Equal Length Equivalent Slope

Do not use a basin-average slope as though it were the slope of a
specific upland/headwater zone.

Do not use one slope metric as a substitute for another.

------------------------------------------------------------

### Drainage Density and Infiltration

Interpret drainage density, stream frequency, constant of channel
maintenance and infiltration number as morphometric indicators.

Do not describe them as direct measurements of:
* soil permeability
* infiltration capacity
* hydraulic conductivity
* groundwater recharge

unless those data are explicitly provided.

============================================================
8. ENGINEERING RECOMMENDATIONS
============================================================

Provide practical watershed-management and engineering recommendations
based on the GIS evidence.

Consider, where appropriate:

* Check dams
* Gully plugs
* Contour bunding
* Contour trenches
* Slope stabilization
* Sediment-control measures
* Drainage infrastructure improvements
* Channel management
* Floodplain management
* Erosion-control measures

Every recommendation must be linked to an observed or reported
catchment characteristic.

Use this structure:

| Intervention | Suitable Terrain/Condition | Purpose | Supporting GIS Evidence |

Do not identify an exact construction site unless the supplied GIS
data are sufficient to justify that site.

If exact site selection requires additional information, state the
additional information required, such as:

* local slope raster
* drainage-order map
* stream-power analysis
* land-use/land-cover
* soil map
* geology
* rainfall
* cross-sections
* hydraulic modelling
* field verification

Do not manufacture site-specific thresholds.

============================================================
9. SEDIMENT AND EROSION INTERPRETATION
============================================================

Where relief, slope, drainage characteristics or terrain patterns
suggest potential erosion concerns, discuss them as potential risks.

Distinguish:

* susceptibility indicator
from
* measured erosion rate.

Do not claim an actual erosion rate unless supplied.

Recommendations for sediment management should be tied to the relevant
terrain/drainage characteristics.

============================================================
10. GIS QA/QC AND CALCULATION NOTES
============================================================

Include important processing notes from the source.

Preserve the source's terminology and calculation-method descriptions.

Discuss, where applicable:

* Pour-point processing
* DEM clipping
* Boundary conditions
* NoData conditions
* Endpoint handling
* Slope calculation methods
* Iterative/equivalent-slope calculations
* DEM resolution/noise
* Any other GIS processing notes

Do not automatically classify implementation-specific status labels as
errors.

Explain what the source reports and why the information may matter
for interpretation.

============================================================
11. REQUIRED REPORT STRUCTURE
============================================================

### 1. Executive Summary

Provide a concise high-level overview of the catchment.

Include:

* Overall basin geometry
* Terrain/elevation characteristics
* Drainage characteristics
* Hydrological response
* Principal engineering implications
* Important GIS/calculation considerations

Include a table:

| Metric | Reported Value | Engineering Significance |

Use exact reported values.

Also include several concise spatial observations derived from the
GIS figures.

------------------------------------------------------------

### 2. Basin Morphometry & Shape Analysis

Discuss:

* Form Factor
* Elongation Ratio
* Circularity Ratio
* Compactness Coefficient
* Basin geometry visible in the GIS figure

Explain implications for runoff concentration, hydrograph shape,
peak-flow timing and travel time.

------------------------------------------------------------

### 3. Hydrological Response & Flood Behaviour

Analyze:

* Time of Concentration
* Basin Relief
* Relief Ratio
* Longest Flow Path
* Longest Flow Path Slope
* Main Channel Length
* Main Channel Slope

Discuss:
* runoff timing
* overland flow
* channelized flow
* basin-scale response
* localized headwater response

Do not make unsupported flood-risk classifications.

------------------------------------------------------------

### 4. Drainage Network Insights

Analyze:

* Stream Order
* Stream Numbers
* Total Stream Length
* Drainage Density
* Stream Frequency
* Constant of Channel Maintenance
* Infiltration Number
* Ruggedness Number
* Bifurcation Ratios
* Stream Length Ratios

Integrate the actual stream-network map.

Discuss:
* drainage texture
* branching pattern
* spatial distribution
* runoff connectivity
* morphometric indications of infiltration/runoff behaviour
* potential erosion/sediment implications

------------------------------------------------------------

### 5. Terrain, Contour & Hypsometric Analysis

Integrate:

* Minimum Elevation
* Maximum Elevation
* Mean Elevation
* Basin Relief
* True Average Basin Slope
* Contour Map
* Hypsometric Curve
* Hypsometric Integral
* Reported Geomorphic Stage

Describe the broad spatial elevation pattern visible in the map.

Interpret the hypsometric curve and its reported metrics.

------------------------------------------------------------

### 6. Longest Flow Path vs. Main Channel

Explicitly compare:

* Longest Flow Path
* Main Channel

Discuss:

* Length
* Starting point
* Outlet/end point
* Spatial route
* Slope characteristics
* Hydrological significance

Clearly explain why the two are distinct GIS/hydrological features.

------------------------------------------------------------

### 7. Engineering & Watershed Management Recommendations

Provide evidence-based recommendations.

For each recommendation include:

| Intervention | Suitable Terrain/Condition | Purpose | Supporting GIS Evidence |

Distinguish between:
* basin-level recommendations
* recommendations requiring detailed site investigation.

------------------------------------------------------------

### 8. GIS QA/QC and Calculation Notes

Summarize relevant GIS processing and calculation information.

Preserve numerical values and reported methods/statuses.

------------------------------------------------------------

### 9. Overall Engineering Conclusion

Provide a concise professional conclusion covering:

1. Basin morphology
2. Terrain and elevation
3. Drainage characteristics
4. Hydrological response
5. Engineering/watershed-management priorities
6. Important GIS/calculation considerations
7. Recommended next steps for detailed investigation

============================================================
12. WRITING AND FORMATTING
============================================================

Use professional engineering and hydrological language.

The report must be specific to the supplied catchment.

Avoid generic textbook explanations unless they directly help interpret
a reported catchment characteristic.

Do not repeatedly restate the same numerical value.

Use simple Markdown:

### Main sections
## Subsections
* Bullet points
Markdown tables where useful

DO NOT use ASCII art or text-based diagrams.

DO NOT use complex LaTeX such as:
\\frac
$$
\\[
\\]

Use simple inline mathematical notation such as:
A / L²

============================================================
13. FINAL QUALITY CONTROL
============================================================

Before producing the final report, verify all of the following:

[ ] Exact numerical values from the source are preserved.

[ ] All supplied GIS images have been visually inspected.

[ ] The report contains substantive image-derived observations.

[ ] Text/tables and maps have been cross-checked.

[ ] Basin shape metrics agree conceptually with the observed geometry.

[ ] Stream statistics are interpreted together with the stream-network
    map.

[ ] Longest Flow Path and Main Channel are clearly distinguished.

[ ] Different slope methods are kept separate.

[ ] Basin-average slope is not incorrectly assigned to specific local
    terrain zones.

[ ] Calculation status labels are not automatically treated as errors.

[ ] Time of Concentration is not used alone to classify flood risk.

[ ] Drainage Density/Infiltration Number are not treated as direct
    measurements of soil permeability.

[ ] No unsupported geological, soil, groundwater, land-use, tidal,
    hydraulic-capacity or waterlogging claims are introduced.

[ ] No exact local slope/elevation thresholds are invented.

[ ] Hypsometric Integral is not converted into an exact percentage of
    rock volume eroded.

[ ] Engineering recommendations are connected to specific GIS evidence.

[ ] Recommendations requiring additional site data clearly identify
    those data requirements.

[ ] GIS processing notes are represented accurately.

[ ] No numerical values have been recalculated.

[ ] No external data has been invented.

Return ONLY the completed professional engineering report.
"""


def extract_text_from_docx(file_path: str) -> str:
    """Reads all paragraphs from a DOCX file and returns them as a single string."""
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

def save_text_to_docx(text: str, output_path: str):
    doc = docx.Document()
    lines = text.split('\n')
    
    in_table = False
    table = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 1. Handle Markdown Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                # Count columns based on the first row
                cols = len([c for c in line.split('|') if c.strip()])
                table = doc.add_table(rows=0, cols=cols)
                table.style = 'Table Grid'
            
            # Skip the alignment row (e.g., |---|---|)
            if '---' in line:
                continue
                
            row_cells = table.add_row().cells
            # Extract cell content and strip bold formatting
            cell_data = [c.strip().replace('**', '') for c in line.split('|') if c.strip()]
            
            for i, val in enumerate(cell_data):
                if i < len(row_cells):
                    row_cells[i].text = val
            continue
        else:
            in_table = False

        # 2. Clean inline Markdown tags (Bold)
        clean_line = line.replace('**', '')

        # 3. Handle Headings
        if line.startswith('### '):
            doc.add_heading(clean_line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(clean_line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(clean_line.replace('# ', ''), level=1)
            
        # 4. Handle Bullets
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(clean_line[2:], style='List Bullet')
            
        # 5. Handle Normal Paragraphs
        else:
            doc.add_paragraph(clean_line)
            
    doc.save(output_path)

# ... [imports and save_text_to_docx function stay the same] ...

def generate_hydrology_report(input_docx_path: str, output_docx_path: str, api_key: str):
    print(f"Uploading DOCX directly: {input_docx_path}")
    
    client = genai.Client(api_key=api_key) 
    
    # 1. Upload the DOCX file directly using the Files API
    doc_file = client.files.upload(file=input_docx_path)

    # Extract embedded images from the DOCX
    doc = docx.Document(input_docx_path)

    image_files = []

    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            image_part = rel.target_part
            image_bytes = image_part.blob

            # Upload image to Gemini
            image_file = client.files.upload(
                file=io.BytesIO(image_bytes),
                config={
                    "mime_type": "image/png"
                }
            )

            image_files.append(image_file)
            print(f"Uploaded image {i + 1}")

    print("Sending DOCX + images to Google Gemini API...")
    
    # 2. Pass both the file object and the prompt string in a list
    response = client.models.generate_content(
        model='gemini-3.6-flash',
        contents=[
            doc_file, 
            *image_files,
            PROMPT_TEMPLATE
        ]
    )
    
    print(f"Saving generated insights to: {output_docx_path}")
    save_text_to_docx(response.text, output_docx_path)
    
    # 3. Clean up the uploaded file from the server
    client.files.delete(name=doc_file.name)
    for image_file in image_files:
        client.files.delete(name=image_file.name)
    print("Report generation complete!")


"""
You are a Senior Hydrologist and GIS Specialist. Your task is to analyze the provided GIS catchment/watershed output and generate a professional engineering insights report.

CRITICAL INSTRUCTIONS:
- You MUST strictly use the exact numerical values provided in the input text. Do not recalculate, estimate, or alter any metrics (e.g., Area, Perimeter, Slopes, Stream Orders).
- DO NOT use ASCII art or text-based diagrams. 
- DO NOT use complex LaTeX formatting (e.g., avoid \\frac or $$ blocks). Use standard inline text for math (e.g., A / L^2) so it renders safely in Microsoft Word.
- Use simple Markdown formatting: '###' for main sections, '##' for sub-sections, and '*' for bullet points.

### Sections to Include in the Report:

### 1. Executive Summary
*   Provide a brief, high-level overview of the catchment based on the data.
*   Create a clean summary table of the most critical metrics (Area, Relief, Longest Flow Path, Main Channel Slope).

### 2. Basin Morphometry & Shape Analysis
*   Extract the Form Factor, Elongation Ratio, and Circularity Ratio exactly as provided.
*   Interpret these specific values to explain how the basin shape will likely affect flood hydrographs, peak discharge, and time to peak.

### 3. Hydrological Response & Flood Behavior
*   Analyze the Time of Concentration (Tc) and Basin Relief.
*   Discuss what these figures mean for runoff timing, overland flow dynamics, and vulnerability to flash floods.

### 4. Drainage Network Insights
*   Analyze the Drainage Density, Stream Frequency, and Constant of Channel Maintenance.
*   Interpret what these metrics suggest about the terrain's infiltration capacity, erodibility, and runoff efficiency.

### 5. Engineering & Watershed Management Recommendations
*   Based on the data, provide actionable engineering interventions. 
*   Specifically recommend suitable locations or conditions for structures like check dams, contour bunding, or drainage infrastructure upgrades. Include notes on sediment control if the erodibility metrics suggest it.
*   Acknowledge any GIS calculation warnings (e.g., pour point snapping or boundary mismatches) and recommend steps to resolve them.

"""